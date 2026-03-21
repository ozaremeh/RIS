# core/src/document_tools.py

from docx import Document
from docx2python import docx2python
from typing import List, Dict, Any


# ------------------------------------------------------------
# Utility: recursively flatten nested lists into strings
# ------------------------------------------------------------

def flatten_runs(item) -> List[str]:
    """Recursively flatten docx2python run structures into a list of strings."""
    if isinstance(item, str):
        return [item]

    if isinstance(item, list):
        flat = []
        for sub in item:
            flat.extend(flatten_runs(sub))
        return flat

    return [str(item)]


# ------------------------------------------------------------
# Load a .docx file (for saving/editing)
# ------------------------------------------------------------

def load_docx(path: str) -> Document:
    return Document(path)


# ------------------------------------------------------------
# Safe table extraction across all docx2python versions
# ------------------------------------------------------------

def extract_tables(doc) -> List[List[List[str]]]:
    """
    Extract tables from docx2python output safely.
    Works across all versions of docx2python.
    Returns a list of tables, each table = list of rows, each row = list of cell strings.
    """

    tables = []

    # Case 1: Newer versions expose doc.tables
    if hasattr(doc, "tables") and isinstance(doc.tables, list):
        for table in doc.tables:
            rows = []
            for row in table:
                row_text = []
                for cell in row:
                    flat = flatten_runs(cell)
                    row_text.append(" ".join(flat).strip())
                rows.append(row_text)
            tables.append(rows)
        return tables

    # Case 2: Some versions embed tables inside doc.body
    # doc.body = [paragraphs_section, tables_section, ...]
    if hasattr(doc, "body") and isinstance(doc.body, list):
        for section in doc.body:
            if isinstance(section, list):
                # A table is a list of rows, where each row is a list of cells
                if section and isinstance(section[0], list) and isinstance(section[0][0], list):
                    rows = []
                    for row in section:
                        row_text = []
                        for cell in row:
                            flat = flatten_runs(cell)
                            row_text.append(" ".join(flat).strip())
                        rows.append(row_text)
                    tables.append(rows)

    return tables


# ------------------------------------------------------------
# Extract document structure using docx2python
# ------------------------------------------------------------

def extract_text_structure(path: str) -> Dict[str, Any]:
    doc = docx2python(path)
    content = []

    # Extract paragraphs
    for section in doc.body:
        for para in section:
            flat = flatten_runs(para)
            text = " ".join(flat).strip()
            if text:
                content.append({"type": "paragraph", "text": text})

    # Extract tables (robust)
    for table in extract_tables(doc):
        content.append({"type": "table", "rows": table})

    return {"content": content}


# ------------------------------------------------------------
# Extract comments
# ------------------------------------------------------------

def extract_comments(path: str) -> List[Dict[str, str]]:
    doc = docx2python(path)
    return [{"text": c} for c in doc.comments]


# ------------------------------------------------------------
# Tracked changes (not exposed by docx2python)
# ------------------------------------------------------------

def extract_tracked_changes(path: str) -> List[Dict[str, str]]:
    return []


# ------------------------------------------------------------
# Combine everything
# ------------------------------------------------------------

def extract_full_document_context(path: str) -> Dict[str, Any]:
    structure = extract_text_structure(path)
    comments = extract_comments(path)
    changes = extract_tracked_changes(path)

    return {
        "structure": structure,
        "comments": comments,
        "tracked_changes": changes,
    }


# ------------------------------------------------------------
# Apply model edits (python-docx)
# ------------------------------------------------------------

def apply_model_edits(doc: Document, edits: List[Dict[str, Any]]) -> Document:
    for edit in edits:
        action = edit.get("action")

        if action == "replace":
            old = edit.get("target", "")
            new = edit.get("new", "")
            for p in doc.paragraphs:
                if old in p.text:
                    p.text = p.text.replace(old, new)

        elif action == "append":
            doc.add_paragraph(edit.get("text", ""))

        elif action == "insert_after_heading":
            heading = edit.get("heading", "")
            text = edit.get("text", "")
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip() == heading:
                    new_p = doc.add_paragraph(text)
                    p._p.addnext(new_p._p)
                    break

    return doc


# ------------------------------------------------------------
# Save document
# ------------------------------------------------------------

def save_docx(doc: Document, path: str) -> None:
    doc.save(path)
