# src/ingestion/extract_text.py

from __future__ import annotations
from pathlib import Path
from typing import Optional
from pypdf import PdfReader
from docx import Document


# ------------------------------------------------------------
# PDF extraction
# ------------------------------------------------------------

def extract_pdf(path: str) -> str:
    """
    Extract text from a PDF using pypdf.
    This handles most scientific papers reliably.
    """
    reader = PdfReader(path)
    pages = []

    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append(text)

    return "\n".join(pages)


# ------------------------------------------------------------
# DOCX extraction
# ------------------------------------------------------------

def extract_docx(path: str) -> str:
    """
    Extract text from a .docx file using python-docx.
    """
    doc = Document(path)
    parts = []

    for para in doc.paragraphs:
        parts.append(para.text)

    # Tables (optional)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            parts.append(row_text)

    return "\n".join(parts)


# ------------------------------------------------------------
# Plain text fallback
# ------------------------------------------------------------

def extract_plain(path: str) -> str:
    """
    Fallback for .txt, .md, etc.
    """
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# ------------------------------------------------------------
# Normalization
# ------------------------------------------------------------

def normalize_text(text: str) -> str:
    """
    Clean up whitespace, normalize line breaks, etc.
    """
    # Replace Windows line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Collapse excessive blank lines
    lines = [line.rstrip() for line in text.split("\n")]
    cleaned = []
    blank_count = 0

    for line in lines:
        if line.strip() == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned.append("")
        else:
            blank_count = 0
            cleaned.append(line)

    return "\n".join(cleaned).strip()


# ------------------------------------------------------------
# Main entry point
# ------------------------------------------------------------

def extract_text(path: str) -> str:
    """
    Extract text from PDF, DOCX, or plain text files.
    """
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        raw = extract_pdf(path)
    elif ext == ".docx":
        raw = extract_docx(path)
    elif ext in {".txt", ".md"}:
        raw = extract_plain(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return normalize_text(raw)
